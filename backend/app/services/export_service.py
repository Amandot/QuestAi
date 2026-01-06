from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import List
from ..models.quiz import Quiz
from ..models.question import Question
import io

class ExportService:
    def __init__(self):
        pass
    
    def generate_exam_paper(self, quiz: Quiz) -> io.BytesIO:
        """Generate a printable exam paper in Word format"""
        doc = Document()
        
        # Add header section for school details
        self._add_header_section(doc, quiz)
        
        # Add instructions
        self._add_instructions(doc)
        
        # Add questions
        self._add_questions_section(doc, quiz.questions)
        
        # Add answer key on separate page
        doc.add_page_break()
        self._add_answer_key(doc, quiz.questions)
        
        # Save to BytesIO
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        
        return doc_io
    
    def _add_header_section(self, doc: Document, quiz: Quiz):
        """Add header section with school details and exam info"""
        # School name placeholder
        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = header.add_run("_" * 50)
        run.bold = True
        
        header.add_run("\nSCHOOL NAME")
        header.add_run("\n" + "_" * 50)
        
        # Exam details
        doc.add_paragraph()  # Empty line
        
        details_table = doc.add_table(rows=4, cols=2)
        details_table.style = 'Table Grid'
        
        # Fill in exam details
        details_table.cell(0, 0).text = "Subject:"
        details_table.cell(0, 1).text = "_" * 30
        
        details_table.cell(1, 0).text = "Class:"
        details_table.cell(1, 1).text = "_" * 30
        
        details_table.cell(2, 0).text = "Time:"
        details_table.cell(2, 1).text = "_" * 30
        
        details_table.cell(3, 0).text = "Max Marks:"
        details_table.cell(3, 1).text = f"{len(quiz.questions)} marks"
        
        # Student details
        doc.add_paragraph()
        student_table = doc.add_table(rows=2, cols=2)
        student_table.style = 'Table Grid'
        
        student_table.cell(0, 0).text = "Name:"
        student_table.cell(0, 1).text = "_" * 40
        
        student_table.cell(1, 0).text = "Roll No:"
        student_table.cell(1, 1).text = "_" * 40
        
        doc.add_paragraph()  # Empty line
    
    def _add_instructions(self, doc: Document):
        """Add general instructions for the exam"""
        instructions_heading = doc.add_paragraph("INSTRUCTIONS:")
        instructions_heading.runs[0].bold = True
        
        instructions = [
            "1. Read all questions carefully before answering.",
            "2. Write your answers in the space provided.",
            "3. For multiple choice questions, circle the correct option.",
            "4. Attempt all questions.",
            "5. Check your answers before submitting."
        ]
        
        for instruction in instructions:
            doc.add_paragraph(instruction)
        
        doc.add_paragraph()  # Empty line
        
        # Add a horizontal line
        doc.add_paragraph("_" * 80)
        doc.add_paragraph()
    
    def _add_questions_section(self, doc: Document, questions: List[Question]):
        """Add all questions with proper formatting and spacing"""
        questions_heading = doc.add_paragraph("QUESTIONS:")
        questions_heading.runs[0].bold = True
        questions_heading.runs[0].underline = True
        
        doc.add_paragraph()
        
        for i, question in enumerate(questions, 1):
            # Question number and text
            q_para = doc.add_paragraph()
            q_para.add_run(f"Q{i}. ").bold = True
            q_para.add_run(question.question_text)
            
            # Add marks indication
            marks_para = doc.add_paragraph()
            marks_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            marks_para.add_run("[1 mark]").italic = True
            
            # Add options for MCQ
            if question.question_type == "MCQ" and question.options:
                for j, option in enumerate(question.options):
                    option_para = doc.add_paragraph()
                    option_para.add_run(f"    {chr(65 + j)}) {option}")
            
            # Add answer space
            if question.question_type == "Short Answer":
                doc.add_paragraph("Answer:")
                for _ in range(3):  # 3 lines for short answer
                    doc.add_paragraph("_" * 70)
            elif question.question_type == "True/False":
                doc.add_paragraph("Answer: True / False (Circle the correct option)")
            
            doc.add_paragraph()  # Space between questions
    
    def _add_answer_key(self, doc: Document, questions: List[Question]):
        """Add answer key on a separate page"""
        answer_key_heading = doc.add_paragraph("ANSWER KEY")
        answer_key_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        answer_key_heading.runs[0].bold = True
        answer_key_heading.runs[0].underline = True
        
        doc.add_paragraph()
        
        for i, question in enumerate(questions, 1):
            answer_para = doc.add_paragraph()
            answer_para.add_run(f"Q{i}. ").bold = True
            
            if question.question_type == "MCQ" and question.options:
                # Find the index of correct answer
                try:
                    correct_index = question.options.index(question.correct_answer)
                    answer_para.add_run(f"{chr(65 + correct_index)}) {question.correct_answer}")
                except ValueError:
                    answer_para.add_run(question.correct_answer)
            else:
                answer_para.add_run(question.correct_answer)
            
            # Add source reference for teachers
            if question.source_page:
                source_para = doc.add_paragraph()
                source_para.add_run(f"    Source: Page {question.source_page}").italic = True
                if question.bloom_level:
                    source_para.add_run(f" | Bloom Level: {question.bloom_level}").italic = True
            
            doc.add_paragraph()  # Space between answers